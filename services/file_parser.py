import os
import io
import filetype  # Replace magic with filetype
import docx
import logging
from werkzeug.utils import secure_filename
from docx.opc.exceptions import PackageNotFoundError
from services.pdf_extractor import PDFExtractor  # Import the new PDFExtractor class

logger = logging.getLogger(__name__)

class FileParser:
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB in bytes
    ALLOWED_EXTENSIONS = {'md', 'docx', 'pdf'}
    ALLOWED_MIMETYPES = {
        'text/plain': 'md',             # For Markdown files
        'text/markdown': 'md',          # Alternative MIME type for Markdown
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/pdf': 'pdf'
    }

    def __init__(self):
        """Initialize FileParser with a PDFExtractor instance"""
        self.pdf_extractor = PDFExtractor()
        logger.info("Initialized FileParser with PDFExtractor")

    @staticmethod
    def allowed_file(file):
        """Check if file type is allowed and size is within limit"""
        try:
            if not file:
                return False, "No file provided"

            # Check file size
            file.seek(0, os.SEEK_END)
            size = file.tell()
            file.seek(0)
            if size > FileParser.MAX_FILE_SIZE:
                return False, "File size exceeds 5MB limit"

            # Get filename and extension
            filename = secure_filename(file.filename)
            extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None

            logger.debug(f"File validation - Name: {filename}, Extension: {extension}")

            if not extension:
                return False, "No file extension found"

            if extension not in FileParser.ALLOWED_EXTENSIONS:
                return False, f"File extension .{extension} not allowed"

            # Check file type using filetype instead of python-magic
            file_content = file.read()
            file.seek(0)
            
            # Special handling for Markdown files (text files)
            if extension == 'md':
                # Simple check for text files - try to decode as UTF-8
                try:
                    file_content.decode('utf-8')
                    mime_type = 'text/plain'
                except UnicodeDecodeError:
                    return False, "Invalid markdown file format"
            else:
                # Use filetype for binary files (PDF, DOCX)
                kind = filetype.guess(file_content)
                if kind is None:
                    return False, "Unknown file type"
                mime_type = kind.mime

            logger.debug(f"File MIME type: {mime_type}")

            if mime_type not in FileParser.ALLOWED_MIMETYPES:
                return False, f"File type {mime_type} not allowed"

            # Verify extension matches MIME type
            if FileParser.ALLOWED_MIMETYPES[mime_type] != extension:
                return False, "File extension doesn't match its content type"

            return True, None

        except Exception as e:
            logger.error(f"Error validating file: {str(e)}")
            return False, f"Error validating file: {str(e)}"

    def parse_to_markdown(self, file):
        """Parse different file types to markdown format"""
        try:
            # Get file info
            filename = secure_filename(file.filename)
            extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None

            logger.debug(f"Parsing file: {filename}")

            # Read file content
            file_content = file.read()
            file.seek(0)
            
            # Determine mime type
            if extension == 'md':
                mime_type = 'text/plain'
            else:
                kind = filetype.guess(file_content)
                mime_type = kind.mime if kind else None

            logger.debug(f"File MIME type for parsing: {mime_type}")

            # Handle Markdown files (both text/plain and text/markdown)
            if extension == 'md':
                content = file.read().decode('utf-8')
                logger.debug(f"Parsed markdown content length: {len(content)}")
                return content

            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = docx.Document(io.BytesIO(file_content))
                markdown = []
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name[-1])
                        markdown.append(f"{'#' * level} {para.text}\n")
                    else:
                        markdown.append(f"{para.text}\n")
                content = '\n'.join(markdown)
                logger.debug(f"Parsed DOCX content length: {len(content)}")
                return content

            elif mime_type == 'application/pdf':
                # Use the simplified PDFExtractor for PDF files
                content = self.pdf_extractor.extract_text(file_content)
                logger.debug(f"Parsed PDF content: {len(content)} chars")
                return content

            else:
                raise ValueError(f"Unsupported file type: {mime_type}")

        except Exception as e:
            logger.error(f"Error parsing file: {str(e)}")
            raise Exception(f"Error parsing file: {str(e)}")

    def parse_file_with_format(self, file):
        """
        Parse file to markdown for display but preserve original format for download
        Returns a tuple of (markdown_content, original_file_bytes, file_format)
        """
        try:
            # Get file info
            filename = secure_filename(file.filename)
            extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None
            
            logger.debug(f"Parsing file with format preservation: {filename}")
            
            # Read file content
            file_content = file.read()
            file.seek(0)
            
            # Determine mime type
            if extension == 'md':
                mime_type = 'text/plain'
            else:
                kind = filetype.guess(file_content)
                mime_type = kind.mime if kind else None
                
            logger.debug(f"File MIME type for format preservation: {mime_type}")
            
            # For Markdown files, just return the content
            if extension == 'md':
                content = file.read().decode('utf-8')
                file.seek(0)
                original_bytes = file.read()
                return content, original_bytes, 'md'
                
            # For DOCX files, convert to markdown for display but keep original bytes
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Parse to markdown for display
                doc = docx.Document(io.BytesIO(file_content))
                markdown = []
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name[-1])
                        markdown.append(f"{'#' * level} {para.text}\n")
                    else:
                        markdown.append(f"{para.text}\n")
                content = '\n'.join(markdown)
                
                # Return markdown content, original bytes, and format
                return content, file_content, 'docx'
                
            # For PDF files, convert to markdown for display but keep original bytes
            elif mime_type == 'application/pdf':
                # Use the simplified PDFExtractor for PDF files
                content = self.pdf_extractor.extract_text(file_content)
                logger.debug(f"Parsed PDF content: {len(content)} chars")
                
                # Return markdown content, original bytes, and format
                return content, file_content, 'pdf'
                
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
                
        except Exception as e:
            logger.error(f"Error parsing file with format preservation: {str(e)}")
            raise Exception(f"Error parsing file: {str(e)}")
            
    @staticmethod
    def markdown_to_docx(markdown_content):
        """
        Convert markdown content back to DOCX format
        Returns the bytes of the DOCX file
        """
        try:
            # Create a new DOCX document
            document = docx.Document()
            
            # Split markdown content by lines
            lines = markdown_content.split('\n')
            
            # Process each line
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Check if it's a heading
                if line.startswith('#'):
                    # Count the number of # to determine heading level
                    level = 0
                    for char in line:
                        if char == '#':
                            level += 1
                        else:
                            break
                            
                    # Add heading with appropriate level
                    if level >= 1 and level <= 9:
                        heading_text = line[level:].strip()
                        document.add_heading(heading_text, level)
                    else:
                        document.add_paragraph(line)
                else:
                    # Regular paragraph
                    document.add_paragraph(line)
                    
            # Save the document to a BytesIO object
            document_io = io.BytesIO()
            document.save(document_io)
            document_io.seek(0)
            
            return document_io.read()
            
        except Exception as e:
            logger.error(f"Error converting markdown to DOCX: {str(e)}")
            raise Exception(f"Error converting markdown to DOCX: {str(e)}")

    @staticmethod
    def markdown_to_pdf(markdown_content):
        """
        Convert markdown content to PDF using ReportLab (with Unicode support)
        
        Args:
            markdown_content: Markdown formatted text
            
        Returns:
            bytes: PDF file content as bytes
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            import io
            import re
            
            # Create a buffer to receive the PDF data
            buffer = io.BytesIO()
            
            # Create the PDF document using ReportLab
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Get the default styles
            styles = getSampleStyleSheet()
            
            # Create custom styles for markdown elements
            styles.add(ParagraphStyle(
                name='CustomHeading1',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12
            ))
            styles.add(ParagraphStyle(
                name='CustomHeading2',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=10
            ))
            styles.add(ParagraphStyle(
                name='CustomHeading3',
                parent=styles['Heading3'],
                fontSize=14,
                spaceAfter=8
            ))
            styles.add(ParagraphStyle(
                name='CustomBulletPoint',
                parent=styles['Normal'],
                fontSize=12,
                leftIndent=20,
                firstLineIndent=0,
                spaceBefore=2,
                spaceAfter=2
            ))
            
            # Process markdown content and build flowables for the document
            flowables = []
            lines = markdown_content.split('\n')
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # Skip empty lines
                if not line:
                    i += 1
                    continue
                
                # Handle headings
                if line.startswith('#'):
                    heading_level = len(re.match(r'^#+', line).group())
                    text = line[heading_level:].strip()
                    
                    if heading_level == 1:
                        flowables.append(Paragraph(text, styles['CustomHeading1']))
                    elif heading_level == 2:
                        flowables.append(Paragraph(text, styles['CustomHeading2']))
                    elif heading_level == 3:
                        flowables.append(Paragraph(text, styles['CustomHeading3']))
                    else:  # h4, h5, h6
                        flowables.append(Paragraph(f"<b>{text}</b>", styles['Normal']))
                    
                    flowables.append(Spacer(1, 0.1 * inch))
                
                # Handle bold text
                elif line.startswith('**') and line.endswith('**'):
                    text = line.strip('**')
                    flowables.append(Paragraph(f"<b>{text}</b>", styles['Normal']))
                
                # Handle list items
                elif line.startswith('- ') or line.startswith('* '):
                    text = '• ' + line[2:]  # Use bullet character
                    flowables.append(Paragraph(text, styles['CustomBulletPoint']))
                
                # Handle numbered list
                elif re.match(r'^\d+\.', line):
                    text = line  # Keep the numbering
                    flowables.append(Paragraph(text, styles['CustomBulletPoint']))
                
                # Regular text
                else:
                    flowables.append(Paragraph(line, styles['Normal']))
                
                i += 1
            
            # Build the PDF document
            doc.build(flowables)
            
            # Get the PDF content from the buffer
            pdf_content = buffer.getvalue()
            buffer.close()
            
            return pdf_content
            
        except Exception as e:
            logger.error(f"Error converting markdown to PDF: {str(e)}")
            raise Exception(f"Error converting markdown to PDF: {str(e)}")